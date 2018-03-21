from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.models import User
from django.utils import timezone
from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponseRedirect, HttpResponse
from django.contrib.auth.decorators import login_required

import os
import speech_recognition as sr
import soundfile as sf
from subprocess import Popen, PIPE, STDOUT
import subprocess
import json
from random import randint
from docx import *
#####################################################
from django.contrib.auth.decorators import login_required
from .models import  Call, Employee
from paralleldots import set_api_key, sentiment, keywords,intent
import pdfkit
from django.utils.translation import ugettext as _
from reportlab.pdfgen import canvas
import reportlab.rl_config
import requests
from twilio.rest import Client
from reportlab.pdfgen import canvas
import pdfcrowd
from gensim.summarization import summarize


#
# Copyright IBM Corp. 2014
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
#
# http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.
#

# Author: Daniel Bolanos
# Date:   2015

# coding=utf-8
import json                        # json
import threading                   # multi threading
import os                          # for listing directories
import queue as Queue                      # queue used for thread syncronization
import sys                         # system calls
import argparse                    # for parsing arguments
import base64                      # necessary to encode in base64
#                                  # according to the RFC2045 standard
import requests                    # python HTTP requests library

# WebSockets
from autobahn.twisted.websocket import WebSocketClientProtocol, \
    WebSocketClientFactory, connectWS
#from twisted.python import log
#from twisted.internet import ssl, reactor

try:
    raw_input          # Python 2
except NameError:
    raw_input = input  # Python 3


class Utils:

    @staticmethod
    def getAuthenticationToken(hostname, serviceName, username, password):

        fmt = hostname + "{0}/authorization/api/v1/token?url={0}/{1}/api"
        uri = fmt.format(hostname, serviceName)
        uri = uri.replace("wss://", "https://").replace("ws://", "https://")
        print(uri)
        auth = (username, password)
        headers = {'Accept': 'application/json'}
        resp = requests.get(uri, auth=auth, verify=False, headers=headers,
                            timeout=(30, 30))
        print(resp.text)
        jsonObject = resp.json()
        return jsonObject['token']


class WSInterfaceFactory(WebSocketClientFactory):

    def __init__(self, queue, summary, dirOutput, contentType, model,
                 url=None, headers=None, debug=None):

        WebSocketClientFactory.__init__(self, url=url, headers=headers)
        self.queue = queue
        self.summary = summary
        self.dirOutput = dirOutput
        self.contentType = contentType
        self.model = model
        self.queueProto = Queue.Queue()

        self.openHandshakeTimeout = 10
        self.closeHandshakeTimeout = 10

        # start the thread that takes care of ending the reactor so
        # the script can finish automatically (without ctrl+c)
        endingThread = threading.Thread(target=self.endReactor, args=())
        endingThread.daemon = True
        endingThread.start()

    def prepareUtterance(self):

        try:
            utt = self.queue.get_nowait()
            self.queueProto.put(utt)
            return True
        except Queue.Empty:
            print("getUtterance: no more utterances to process, queue is "
                  "empty!")
            return False

    def endReactor(self):

        self.queue.join()
        print("about to stop the reactor!")
        reactor.stop()

    # this function gets called every time connectWS is called (once
    # per WebSocket connection/session)
    def buildProtocol(self, addr):

        try:
            utt = self.queueProto.get_nowait()
            proto = WSInterfaceProtocol(self, self.queue, self.summary,
                                        self.dirOutput, self.contentType)
            proto.setUtterance(utt)
            return proto
        except Queue.Empty:
            print("queue should not be empty, otherwise this function should "
                  "not have been called")
            return None


# WebSockets interface to the STT service
#
# note: an object of this class is created for each WebSocket
# connection, every time we call connectWS
class WSInterfaceProtocol(WebSocketClientProtocol):

    def __init__(self, factory, queue, summary, dirOutput, contentType):
        self.factory = factory
        self.queue = queue
        self.summary = summary
        self.dirOutput = dirOutput
        self.contentType = contentType
        self.packetRate = 20
        self.listeningMessages = 0
        self.timeFirstInterim = -1
        self.bytesSent = 0
        self.chunkSize = 2000     # in bytes
        super(self.__class__, self).__init__()
        print(dirOutput)
        print("contentType: {} queueSize: {}".format(self.contentType,
                                                     self.queue.qsize()))

    def setUtterance(self, utt):

        self.uttNumber = utt[0]
        self.uttFilename = utt[1]
        self.summary[self.uttNumber] = {"hypothesis": "",
                                        "status": {"code": "", "reason": ""}}
        self.fileJson = "{}/{}.json.txt".format(self.dirOutput, self.uttNumber)
        try:
            os.remove(self.fileJson)
        except OSError:
            pass

    # helper method that sends a chunk of audio if needed (as required
    # what the specified pacing is)
    def maybeSendChunk(self, data):

        def sendChunk(chunk, final=False):
            self.bytesSent += len(chunk)
            self.sendMessage(chunk, isBinary=True)
            if final:
                self.sendMessage(b'', isBinary=True)

        if (self.bytesSent + self.chunkSize >= len(data)):
            if (len(data) > self.bytesSent):
                sendChunk(data[self.bytesSent:len(data)], True)
                return
        sendChunk(data[self.bytesSent:self.bytesSent + self.chunkSize])
        self.factory.reactor.callLater(0.01, self.maybeSendChunk, data=data)
        return

    def onConnect(self, response):
        print("onConnect, server connected: {}".format(response.peer))

    def onOpen(self):
        print("onOpen")
        data = {"action": "start",
                "content-type": str(self.contentType),
                "continuous": True,
                "interim_results": True,
                "inactivity_timeout": 600,
                'max_alternatives': 3,
                'timestamps': True,
                'word_confidence': True}
        print("sendMessage(init)")
        # send the initialization parameters
        self.sendMessage(json.dumps(data).encode('utf8'))

        # start sending audio right away (it will get buffered in the
        # STT service)
        print(self.uttFilename)
        with open(str(self.uttFilename), 'rb') as f:
            self.bytesSent = 0
            dataFile = f.read()
        self.maybeSendChunk(dataFile)
        print("onOpen ends")

    def onMessage(self, payload, isBinary):

        if isBinary:
            print("Binary message received: {0} bytes".format(len(payload)))
        else:
            print(u"Text message received: {0}".format(payload.decode('utf8')))

            # if uninitialized, receive the initialization response
            # from the server
            jsonObject = json.loads(payload.decode('utf8'))
            if 'state' in jsonObject:
                self.listeningMessages += 1
                if self.listeningMessages == 2:
                    print("sending close 1000")
                    # close the connection
                    self.sendClose(1000)

            # if in streaming
            elif 'results' in jsonObject:
                jsonObject = json.loads(payload.decode('utf8'))
                hypothesis = ""
                # empty hypothesis
                if len(jsonObject['results']) == 0:
                    print("empty hypothesis!")
                # regular hypothesis
                else:
                    # dump the message to the output directory
                    jsonObject = json.loads(payload.decode('utf8'))
                    with open(self.fileJson, "a") as f:
                        f.write(json.dumps(jsonObject, indent=4,
                                           sort_keys=True))

                    res = jsonObject['results'][0]
                    hypothesis = res['alternatives'][0]['transcript']
                    bFinal = (res['final'] is True)
                    if bFinal:
                        print('final hypothesis: "' + hypothesis + '"')
                        self.summary[self.uttNumber]['hypothesis'] += hypothesis
                    else:
                        print('interim hyp: "' + hypothesis + '"')

    def onClose(self, wasClean, code, reason):

        print("onClose")
        print("WebSocket connection closed: {0}, code: {1}, clean: {2}, "
              "reason: {0}".format(reason, code, wasClean))
        self.summary[self.uttNumber]['status']['code'] = code
        self.summary[self.uttNumber]['status']['reason'] = reason

        # create a new WebSocket connection if there are still
        # utterances in the queue that need to be processed
        self.queue.task_done()

        if not self.factory.prepareUtterance():
            return

        # SSL client context: default
        if self.factory.isSecure:
            contextFactory = ssl.ClientContextFactory()
        else:
            contextFactory = None
        connectWS(self.factory, contextFactory)


def tryfunc():

    with open('recordings.txt', 'w') as f:
        rootdir = os.getcwd() + '/recordings/'
        for subdir, dirs, files in os.walk(rootdir):
            subdirName = subdir[subdir.rfind('/')+1:]
            # print "hi"
            # print subdirName
            for file in files:
                # print "yo"
                # print file
                f.write('./recordings/' + subdirName + '/' + file)
                f.write('\n')

    args = {'contentType':'audio/wav', 'credentials':['45189e29-8725-4df0-8f94-0599fb8564e9', '1FlfCoi0F7tt'], 'dirOutput':'./output', 'fileInput':'./recordings.txt', 'model':'en-GB_NarrowbandModel', 'optOut':False, 'threads':10, 'tokenauth':False}

    

    # create output directory if necessary
    # if os.path.isdir(args['dirOutput']):
    #     fmt = 'the output directory "{}" already exists, overwrite? (y/n)? '
    #     while True:
    #         answer = raw_input(fmt.format(args['dirOutput'])).strip().lower()
    #         if answer == "n":
    #             sys.stderr.write("exiting...")
    #             sys.exit()
    #         elif answer == "y":
    #             break
    # else:
    #     os.makedirs(args['dirOutput'])

    if not os.path.isdir(args['dirOutput']):
        os.makedirs(args['dirOutput'])


    # logging
    log.startLogging(sys.stdout)

    # add audio files to the processing queue
    q = Queue.Queue()
    lines = [line.rstrip('\n') for line in open(args['fileInput'])]
    fileNumber = 0
    finalDictionary = {'0001': [], '0002': [], '0003': [], '0004': [], '0021': [], '0022':[], '0023':[], '0045':[]}
    myDictionary = {}
    # finalDictionary = {}
    for fileName in lines:
        print(fileName)
        exceptFileName = fileName[:fileName.rfind('/')]
        # print "start"
        # print exceptFileName
        dirName = exceptFileName[exceptFileName.rfind('/')+1:]
        # myDictionary[dirName] = myDictionary[dirName].append(fileNumber)
        # print dirName
        # print fileName
        myDictionary[fileNumber] = dirName
        q.put((fileNumber, fileName))
        fileNumber += 1

    hostname = "stream.watsonplatform.net"
    headers = {'X-WDC-PL-OPT-OUT': '1'} if args['optOut'] else {}

    # authentication header
    if args['tokenauth']:
        headers['X-Watson-Authorization-Token'] = (
            Utils.getAuthenticationToken('https://' + hostname,
                                         'speech-to-text',
                                         '45189e29-8725-4df0-8f94-0599fb8564e9',
                                         '1FlfCoi0F7tt'))
    else:
        auth = args['credentials'][0] + ":" + args['credentials'][1]
        data_bytes = auth.encode("utf-8")
        headers["Authorization"] = "Basic " + base64.b64encode(data_bytes).decode('ascii')

    print(headers)
    # create a WS server factory with our protocol
    fmt = "wss://{}/speech-to-text/api/v1/recognize?model={}"
    url = fmt.format(hostname, args['model'])
    summary = {}
    factory = WSInterfaceFactory(q, summary, args['dirOutput'], args['contentType'],
                                 args['model'], url, headers, debug=False)
    factory.protocol = WSInterfaceProtocol

    for i in range(min(int(args['threads']), q.qsize())):

        factory.prepareUtterance()

        # SSL client context: default
        if factory.isSecure:
            contextFactory = ssl.ClientContextFactory()
        else:
            contextFactory = None
        connectWS(factory, contextFactory)

    reactor.run()

    # dump the hypotheses to the output file
    fileHypotheses = args['dirOutput'] + "/hypotheses.txt"
    f = open(fileHypotheses, "w")
    successful = 0
    emptyHypotheses = 0
    # print sorted(summary.items())
    counter = 0
    for key, value in enumerate(sorted(summary.items())):
        value = value[1]  
        if value['status']['code'] == 1000:
            print('{}: {} {}'.format(key, value['status']['code'],
                                     value['hypothesis'].encode('utf-8')))
            # print "I am a key "
            direc = myDictionary[key] 
            # print value['hypothesis']
            finalDictionary[direc].append(value['hypothesis'])
            # print(fileName)
            # print "YELLOW"
            successful += 1
            if value['hypothesis'][0] == "":
                emptyHypotheses += 1
        else:
            fmt = '{}: {status[code]} REASON: {status[reason]}'
            print(fmt.format(key, **status))
        f.write('{}: {}\n'.format(counter + 1, value['hypothesis'].encode('utf-8')))
        counter += 1
    f.close()
    fmt = "successful sessions: {} ({} errors) ({} empty hypotheses)"
    print(fmt.format(successful, len(summary) - successful, emptyHypotheses))
    # print "YAYYYY"
    return (finalDictionary)





id=0
set_api_key("WjlcgNCnULlVMiRc47ob2ybV0aVS0aR8VhoQUBpayBs")

def getSentiment(callText):
    data = sentiment(callText)
    ans = data['sentiment']
    print(ans)
    return (ans)

def getTextSummary(callText):
    return (summarize(callText))

def getKeywords(callText):
    l=[]
    data = keywords(callText)
    ans = data['keywords']
    for i in ans:
        try:
            k = i.get("keyword")
        except AttributeError:
            l.append(['no keyword'])
            return(l)
    l.append(k)
    print(l)
    return (l)

def getServiceProvider(callText):
    sp=""
    if("Jio" in callText) or ("jio" in callText):
        sp = "jio"
    elif("Airtel" in callText) or ("airtel" in callText):
        sp="airtel"
    elif("Vodafone" in callText) or ("vodafone" in callText):
        sp="vodafone"
    elif("idea" in callText) or ("Idea" in callText):
        sp="idea"
    else:
        sp=""
    return (sp)

def getIntent(callText):
    datater = intent(callText)
    answer = datater['intent']
    print(answer)
    return (answer)

# def insert(request,c):
#     c.save()

def insertCall(callerId,callText,duration):
    sentiment = getSentiment(callText)
    print("SENTIMENT MILA")
    #if callText
    #summary = getTextSummary(callText)
    #print("SUMMARY MILA")
    keywords = getKeywords(callText)
    print ("KEYWORDS ON")
    sp = getServiceProvider(callText)
    print("sp mila")
    intent = getIntent(callText)
    print("INTENT BHI BRO")
    rating = 0 ####FOR NOWWWW
    global id
    id = id+1
    c = Call()
    c.sentiment = sentiment
    c.keywords = keywords
    c.service_provider = sp
    c.intent = intent
    c.rating = rating
    c.ccid = callerId
    c.text = callText
    durationInt = int(float(duration))
    c.duration = durationInt
    #c = call(sentiment=sentiment, keywords=keywords, service_provider=sp, intent=intent, rating=rating, ccid=callerId, text=callText, duration = duration)
    c.save()
    #insert(c)
    print("ADDDDDDDED.")

def getTotalTime(callerId):
    total_time = 0
    calls = Call.objects.filter(ccid = callerId)
    for i in calls:
        total_time = total_time + i.duration
    return (total_time)

def getTotalCalls(callerId):
    #calls = Call.objects.filter(ccid = callerId).values()
    callcount = Call.objects.filter(ccid = callerId).count()
    return (callcount)

def getOverallSentiment(callerId):
    s = 0
    calls = Call.objects.filter(ccid = callerId)
    for i in calls:
        if(i.sentiment=="positive"):
            s = s + 1
        elif(i.sentiment=="negative"):
            s = s - 1
    if (s > 0):
        sentiment = "positive"
    elif(s < 0):
        sentiment = "negative"
    else:
        sentiment = "neutral"
    return (sentiment)


def insertEmployee(callerId):
    total_time = 0
    total = getTotalCalls(callerId)
    total_time = getTotalTime(callerId)
    sentiment = getOverallSentiment(callerId)
    e = Employee()
    e.name = "Purav"
    e.total_calls=total
    e.total_seconds=total_time    
    e.sentiment=sentiment
    #c = call(sentiment=sentiment, keywords=keywords, service_provider=sp, intent=intent, rating=rating, ccid=callerId, text=callText, duration = duration)
    e.save()


# def speech_to_text():
#     finalDictionary = {'0001': [], '0002': [], '0003': [], '0004': []}
#     print ("beginning")
#     with open('recordings.txt', 'w') as f:
#         print ("in here")
#         rootdir = os.getcwd() + '/app/recordings (copy)/'
#         print (rootdir)
#         for subdir, dirs, files in os.walk(rootdir):
#             subdirName = subdir[subdir.rfind('/')+1:]
#             print (subdir)
#             for file in files:
#                 print (file)
#                 # subdirName + '/' + file)
#                 # f.write('\n')
#                 AUDIO_FILE = subdir + '/' + file
#                 r = sr.Recognizer()
#                 with sr.AudioFile(AUDIO_FILE) as source:
#                     print ("hi")
#                     audio = r.record(source) 
#                     IBM_USERNAME = "45189e29-8725-4df0-8f94-0599fb8564e9"  # IBM Speech to Text usernames are strings of the form XXXXXXXX-XXXX-XXXX-XXXX-XXXXXXXXXXXX
#                     IBM_PASSWORD = "1FlfCoi0F7tt"  # IBM Speech to Text passwords are mixed-case alphanumeric strings
#                     try:
#                         f = sf.SoundFile(AUDIO_FILE)
#                         seconds = format(len(f) / f.samplerate)
#                         print (seconds)
#                         print (r.recognize_google(audio))
                        
#                         finalDictionary[subdirName].append((r.recognize_google(audio), seconds))
#                         # print("IBM Speech to Text results:")
#                         # print(r.recognize_ibm(audio, username=IBM_USERNAME, password=IBM_PASSWORD, show_all=True)['results'][0]['alternatives'][0]['transcript'])  # pretty-print the recognition result
#                         # finalDictionary[subdirName].append(r.recognize_ibm(audio, username=IBM_USERNAME, password=IBM_PASSWORD, show_all=True)['results'][0]['alternatives'][0]['transcript'])
#                     except sr.UnknownValueError:
#                         print("IBM Speech to Text could not understand audio")
#                     except sr.RequestError as e:
#                         print("Could not request results from IBM Speech to Text service; {0}".format(e))
    
#     for i in finalDictionary.keys():
#         listOfCalls = finalDictionary.get(i)
#         print(listOfCalls)
#         for j, k in listOfCalls:
#             insertCall(i,j,k)
#             print("starting one insertion")
#         insertEmployee(i)
#     print (finalDictionary)
#     return (finalDictionary)

def retrieve(request):
    query_results = Call.objects.all()
    print(query_results)
    return (query_results)



def dashboard(request):
    return render(request, 'index.html')

def callerlist(request):
    query_results = Employee.objects.all()

    print(query_results)
    return render(request, 'caller-list.html',{'query_results':query_results})

def summary(request):
    return render(request, 'download-summary.html')

def upload(request):
    if request.method == 'POST':
        # dict = speech_to_text()
        # dict = tryfunc()
        # p = ('python3 sttClient.py', shell=True, stdin=PIPE, stdout=PIPE, stderr=STDOUT)
        
        # subprocess.call('python3 sttClient.py', shell = True)
        # stdoutdata, stderrdata = subprocess.Popen.communicate(['python3', 'sttClient.py'])
        # print (stdoutdata)
        # print (output)
        p=subprocess.Popen('python3 sttClient.py',stdout=subprocess.PIPE,shell=True)
        (output,err)=p.communicate()
        p_status=p.wait()
        print (output)
        print (type(output))
        # outputString = output.decode(encoding='UTF-8')
        outputString = str(output)
        
        finalOutputString =  outputString.replace("\\'", "'")
        finalOutputString =  finalOutputString.replace("'", "\"")
        finalOutputString = finalOutputString[finalOutputString.find("{"): finalOutputString.find("}")+1]
        print (finalOutputString)
        mydict = json.loads(finalOutputString)

        for i in mydict.keys():
            listOfCalls = mydict.get(i)
            print(listOfCalls)
            for j in listOfCalls:
                k = randint(15, 35)
                insertCall(i,j,k)
                print("starting one insertion")
        insertEmployee(i)

        return render(request, 'upload.html', {'data': sorted(mydict.items())})
    else:
        print ("Sorry")
        return render(request, 'upload.html')

# def search(request):
#     if request.method == 'POST':
#         search_id = request.POST.get('textfield', None)
#         print (search_id)
#         dict = youtube_search(search_id)
#     return render(request, 'youtube.html', {'data': sorted(dict.items())})

def perid(request):
    callperid = request.GET.get('id')
    calls = Call.objects.filter(ccid = callperid)
    print(calls)
    print("I AM HEREEEEEEEEE")
    return render(request, 'calls-per-id.html',{'calls':calls})

def docDownload(request):
    doc = Document()
    doc.add_paragraph('Summary')
    eid = request.GET.get('eid')
    callObj = Call.objects.filter(ccid=eid)
    callCount = callObj.count()
    table = doc.add_table(rows = (callCount+1), cols = 6)
    row = 0
    table.cell(row,0).text = 'Sentiment'
    table.cell(row,1).text = 'Keywords'
    table.cell(row,2).text = 'Intent'
    table.cell(row,3).text = 'Call Text'
    table.cell(row,4).text = 'Date'
    table.cell(row,5).text = 'Duration'
    row += 1
    duration = 0
    sentimentNeg = 0
    sentimentPos = 0
    sentimentNeu = 0
    for obj in callObj:
        table.cell(row,0).text = obj.sentiment
        table.cell(row,1).text = obj.keywords
        table.cell(row,2).text = obj.intent
        table.cell(row,3).text = obj.text
        table.cell(row,4).text = str(obj.date)
        table.cell(row,5).text = str(obj.duration)
        duration = duration + obj.duration
        if obj.sentiment == "positive":
            sentimentPos = sentimentPos + 1
        elif obj.sentiment == "negative":
            sentimentNeg = sentimentNeg + 1
        elif obj.sentiment == "neutral":
            sentimentNeu = sentimentNeu + 1
        row += 1
    if(sentimentPos>(sentimentNeg+sentimentNeu)):
        status = "Excellent work"
    elif(sentimentNeg>5):
        status = "Warning needed to be given"
    else:
        status = "Consistent"
    response = HttpResponse(content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response['Content-Disposition'] = 'attachment; filename=summary.docx'
    doc.add_paragraph("Thus the person with id " + eid + " has the following status: " + status)
    doc.save(response)
    return response


# Create your views here.
def logout_blog(request):
    print ("hi")
    if request.user.is_authenticated:
        logout(request)
        return render(request,'page-login.html')
    else:
        return HttpResponseRedirect('/login/')

def register(request):
    if request.method == 'POST':
        username = request.POST.get('email')
        password = request.POST.get('password')
        name = request.POST.get('name')
        user = User.objects.create(
            first_name = name,
            username = username,
            )
        user.set_password(password)
        user.save()

        user = authenticate(username = username, password = password)
        login(request, user)
        return redirect('/dashboard/')
    else:
        return render(request,'page-register.html')   

def login_blog(request):
    print ("hiiiiiHDBEHVF")
    if request.method == 'POST':
        username = request.POST.get('email')
        password = request.POST.get('password')
        print (username)
        print (password)
        user = authenticate(username = username, password = password)
        if user:
            print ("hi")
            if user.is_active:
                login(request,user)
                return redirect('/dashboard/')
            else:
                return HttpResponse('Disabled Account')
        else:
            return HttpResponse("Invalid Login details.Are you trying to Sign up?")
    else:
        return render(request,'page-login.html')

############################################
def trial(request,vid):
    call=Call.objects.filter(ccid=vid)
    return render(request, 'trial.html')


def tryIntent(request):
    if request.method == 'POST':
        sentence = request.POST.get('sent')
        datater = intent(sentence)
        answer = datater['intent']
        print(answer)
    return render(request,'trial.html',{'answer':answer})

def index(request):
    pdf = pdfkit.from_url("http://ourcodeworld.com", "ourcodeworld.pdf")

    return HttpResponse("Everything working good, check out the root of your project to see the generated PDF.")

def trySentiment(request):
    if request.method == 'POST':
        sentence = request.POST.get('sent')
        data = sentiment(sentence)
        ans = data['sentiment']
        print(ans)
    return render(request,'trial.html',{'ans':ans})

def call(request):
    mess="test"
    if request.method == 'POST':
        # Your Account SID from twilio.com/console
        account_sid = "ACfca6659db07ae82fa33cafa753466453"
        # Your Auth Token from twilio.com/console
        auth_token  = "f46aaf2f738396b54a5a6c14de0d59a6"

        client = Client(account_sid, auth_token)

        message = client.messages.create(
        to="+919820370451", 
        from_="+18707264156",
        body="Hello niti give me sandwich!")
        

        print(message.sid)
        mess=message.sid
        #data = request.POST.get('sent')
    else:
        print("post not sent")
        print(request.method)
    return render(request,'face.html',{'mess':mess})


def generate_pdf_view(request):
    try:
        # create an API client instance
        client = pdfcrowd.Client("niti", "896b4517ed9216d9932c6147a74fd3ba")

        # convert a web page and store the generated PDF to a variable
        pdf = client.convertURI("http://www.google.com")
        #pdf = client.convertFile("base.html")
         # set HTTP response headers
        response = HttpResponse(content_type="application/pdf")
        response["Cache-Control"] = "max-age=0"
        response["Accept-Ranges"] = "none"
        response["Content-Disposition"] = "attachment; filename=google_com.pdf"

        # send the generated PDF
        response.write(pdf)
    except pdfcrowd.Error:
        response = HttpResponse(mimetype="text/plain")
        response.write(why)
    return response

def pdf_view(request):
    # create an API client instance
    client = pdfcrowd.Client("niti", "896b4517ed9216d9932c6147a74fd3ba")

    # convert a web page and store the generated PDF to a variable
    #pdf = client.convertURI("http://www.google.com")
    pdf = client.convertFile(r"C:\Users\spgna\Documents\Django project\askmeout\app\templates\base.html")
     # set HTTP response headers
    response = HttpResponse(content_type="application/pdf")
    response["Cache-Control"] = "max-age=0"
    response["Accept-Ranges"] = "none"
    response["Content-Disposition"] = "attachment; filename=google_com.pdf"

    # send the generated PDF
    response.write(pdf)
    return response



# def pdf(request):
#     reportlab.rl_config.warnOnMissingFontGlyphs = 0
#     c = canvas.Canvas("./hello.pdf",)
#     c.drawString(100, 100, "Hello World")
#     c.showPage()
#     c.save()
#     print("done")
#     return render(request,'pages.html')

# def pdf_view(request):
#     # create an API client instance
#     client = pdfcrowd.Client("niti", "896b4517ed9216d9932c6147a74fd3ba")

#     # convert a web page and store the generated PDF to a variable
#     #pdf = client.convertURI("http://www.google.com")
#     pdf = client.convertFile(r"C:\Users\spgna\Documents\Django project\askmeout\app\templates\base.html")
#      # set HTTP response headers
#     response = HttpResponse(content_type="application/pdf")
#     response["Cache-Control"] = "max-age=0"
#     response["Accept-Ranges"] = "none"
#     response["Content-Disposition"] = "attachment; filename=google_com.pdf"

    # send the generated PDF
    # response.write(pdf)
    # return response
